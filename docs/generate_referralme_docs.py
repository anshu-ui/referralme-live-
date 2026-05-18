from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
ASSETS = {
    "logo": ROOT / "client" / "public" / "logo.png",
    "reward_shirt": ROOT / "client" / "public" / "reward-shirt.png",
    "reward_kit": ROOT / "client" / "public" / "reward-kit.png",
    "reward_certificate": ROOT / "client" / "public" / "reward-certificate.png",
}


def _set_run_font(run, name: str, size_pt: int | None = None, rgb: tuple[int, int, int] | None = None, bold: bool | None = None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if rgb is not None:
        run.font.color.rgb = RGBColor(*rgb)
    if bold is not None:
        run.bold = bold


def apply_standard_business_brief_styles(doc: Document):
    # Based on documents preset: standard_business_brief
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal_para = normal.paragraph_format
    normal_para.space_before = Pt(0)
    normal_para.space_after = Pt(6)
    normal_para.line_spacing = 1.10

    title = styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(24)
    title.font.bold = True
    title.font.color.rgb = RGBColor(11, 37, 69)  # ink_blue-ish
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(10)
    title.paragraph_format.line_spacing = 1.05

    h1 = styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(16)
    h1.font.color.rgb = RGBColor(46, 116, 181)  # #2E74B5
    h1.paragraph_format.space_before = Pt(16)
    h1.paragraph_format.space_after = Pt(8)
    h1.paragraph_format.line_spacing = 1.05

    h2 = styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(13)
    h2.font.color.rgb = RGBColor(46, 116, 181)
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.line_spacing = 1.05

    h3 = styles["Heading 3"]
    h3.font.name = "Calibri"
    h3.font.size = Pt(12)
    h3.font.color.rgb = RGBColor(31, 77, 120)  # #1F4D78
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after = Pt(4)
    h3.paragraph_format.line_spacing = 1.05


def add_logo_block(doc: Document, subtitle: str):
    if ASSETS["logo"].exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run()
        r.add_picture(str(ASSETS["logo"]), width=Inches(1.35))
    t = doc.add_paragraph("ReferralMe", style="Title")
    t.paragraph_format.space_after = Pt(0)
    sub = doc.add_paragraph(subtitle)
    sub.paragraph_format.space_before = Pt(4)
    sub.paragraph_format.space_after = Pt(12)
    sub_run = sub.runs[0]
    _set_run_font(sub_run, "Calibri", 12, (55, 65, 81), False)


def add_meta_row(doc: Document, items: list[tuple[str, str]]):
    table = doc.add_table(rows=1, cols=len(items))
    table.autofit = True
    row = table.rows[0]
    for idx, (k, v) in enumerate(items):
        cell = row.cells[idx]
        cell.text = ""
        pk = cell.paragraphs[0]
        rk = pk.add_run(f"{k}: ")
        _set_run_font(rk, "Calibri", 10, (31, 41, 55), True)
        rv = pk.add_run(v)
        _set_run_font(rv, "Calibri", 10, (55, 65, 81), False)
    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]):
    for it in items:
        p = doc.add_paragraph(it, style="List Bullet")
        p.paragraph_format.space_after = Pt(4)


def add_numbered(doc: Document, items: list[str]):
    for it in items:
        p = doc.add_paragraph(it, style="List Number")
        p.paragraph_format.space_after = Pt(4)


def product_overview_doc(out_path: Path):
    doc = Document()
    apply_standard_business_brief_styles(doc)

    add_logo_block(doc, "Product Overview and Operating Model")
    add_meta_row(
        doc,
        [
            ("Version", "v1.0"),
            ("Date", date.today().isoformat()),
            ("Use", "Internal"),
        ],
    )

    doc.add_heading("Executive Summary", level=1)
    add_bullets(
        doc,
        [
            "ReferralMe is a career-outcomes platform that helps candidates move from confusion to a clear weekly plan, stronger applications, and verified referral support.",
            "We combine: (1) structured job discovery + referral requests, (2) ATS and application quality tools, (3) AI mentor guidance, and (4) optional paid human mentorship sessions.",
            "A campus ambassador program creates distribution in colleges through weekly tasks, proof review, and reward ladders.",
            "Phase 1 monetization: paid mentorship sessions collected by the platform; subscriptions can be added once the guidance loop proves retention.",
            "North star: increase interview rate and time-to-first-interview for high-intent seekers.",
        ],
    )

    doc.add_heading("Problem", level=1)
    doc.add_paragraph(
        "Most job seekers lack a repeatable system. They send low-signal cold DMs for referrals, apply widely without tailoring, and do not track what works. The result is low response rates and slow progress.",
    )
    add_bullets(
        doc,
        [
            "Referral outreach is noisy and often ignored because it is generic and high volume.",
            "Candidates struggle to identify the right roles, keywords, and projects to match a target job.",
            "Interview preparation is unstructured; effort doesn’t compound week to week.",
        ],
    )

    doc.add_heading("Solution: A Closed-Loop System", level=1)
    add_numbered(
        doc,
        [
            "Intake: understand the target role, experience, constraints, and current status.",
            "Plan: generate a weekly plan (7-day) with tasks, outreach templates, and checkpoints.",
            "Execute: apply to a curated set of roles; send ethical, high-signal referral requests; track outcomes.",
            "Improve: ATS fixes, portfolio upgrades, and interview practice based on feedback loops.",
            "Escalate: when needed, book a paid human mentorship session for mock interviews, resume review, or negotiation.",
        ],
    )

    doc.add_heading("Core Product Modules", level=1)

    doc.add_heading("1) Referral Marketplace (Seeker ↔ Referrer)", level=2)
    add_bullets(
        doc,
        [
            "Seekers: browse roles, request referrals, attach resume/cover letter, and track status.",
            "Referrers: post internal opportunities, review requests, and update status (accepted/rejected/sent-to-HR).",
            "Quality guardrails: discourage “pay-for-referral” language; promote ethical outreach and readiness.",
        ],
    )

    doc.add_heading("2) Tools: ATS + Application Quality", level=2)
    add_bullets(
        doc,
        [
            "ATS score and improvement suggestions based on resume text vs. job description.",
            "Job detail extraction from referral posts/JDs to reduce manual effort for referrers.",
            "Reusable templates for faster, higher-quality applications.",
        ],
    )

    doc.add_heading("3) Mentorship: AI + Human Sessions", level=2)
    add_bullets(
        doc,
        [
            "AI Mentor: guided intake wizard, 7-day plan generation, and chat-based support.",
            "Human mentorship: referrers can enable mentorship, define services (title/duration/price), and confirm sessions with a meeting link.",
            "Payments (Phase 1): platform collects 100% payment at booking; session request is created only after server-side verification.",
            "Payouts (Phase 2): pay out mentors after completion (e.g., weekly) and take a platform fee.",
        ],
    )

    doc.add_heading("4) Campus Ambassador Program", level=2)
    add_bullets(
        doc,
        [
            "Landing page + application form + admin review pipeline.",
            "Ambassador dashboard: tasks, proof submission, activity, points, reward ladder.",
            "Admin: create tasks, review proof, publish announcements and landing content, manage ambassadors.",
        ],
    )

    doc.add_heading("Operating Model (How It Works)", level=1)
    doc.add_heading("Seeker Flow", level=2)
    add_numbered(
        doc,
        [
            "Complete profile and set target role.",
            "Use AI Mentor to generate a weekly plan + outreach scripts.",
            "Apply to roles and submit referral requests with tailored resume.",
            "Track responses and improve with ATS insights.",
            "If stuck, book a mentor session (resume review/mock interview).",
        ],
    )
    doc.add_heading("Referrer Flow", level=2)
    add_numbered(
        doc,
        [
            "Complete profile, post roles, and review incoming requests.",
            "Optionally enable mentorship and publish services/prices.",
            "Confirm booked sessions by adding meeting link; mark sessions complete after delivery.",
        ],
    )

    doc.add_heading("Business Model (Initial)", level=1)
    add_bullets(
        doc,
        [
            "Mentorship: paid sessions (platform collects payment).",
            "Subscription (later): bundled AI guidance + ATS + structured weekly system.",
            "Campus program: distribution channel; can add sponsor partnerships later.",
        ],
    )

    doc.add_heading("Roadmap (Next 90 Days)", level=1)
    add_bullets(
        doc,
        [
            "Stabilize AI mentor with caching + paid-tier quota or provider fallback; keep offline mode as a safety net.",
            "Add mentorship payout ledger and admin payout queue (Phase 2).",
            "Add mentor discovery filters (company, role, rating) and session feedback loop.",
            "Improve referral quality metrics: readiness checklist before requesting referral.",
            "Add analytics dashboard for conversion, interview rate, and retention.",
        ],
    )

    doc.add_heading("Brand Assets (Examples)", level=1)
    doc.add_paragraph("Below are example reward assets used in the campus program and marketing pages.")
    for key, caption in [
        ("reward_shirt", "Reward: Ambassador tee (example mock)."),
        ("reward_kit", "Reward: Brand kit (example mock)."),
        ("reward_certificate", "Reward: Recognition certificate (example mock)."),
    ]:
        p = doc.add_paragraph()
        if ASSETS[key].exists():
            r = p.add_run()
            r.add_picture(str(ASSETS[key]), width=Inches(5.8))
        cap = doc.add_paragraph(caption)
        cap.paragraph_format.space_before = Pt(2)
        cap.paragraph_format.space_after = Pt(10)
        cap_run = cap.runs[0]
        _set_run_font(cap_run, "Calibri", 9, (100, 116, 139), False)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def investor_pitch_doc(out_path: Path):
    doc = Document()
    apply_standard_business_brief_styles(doc)

    add_logo_block(doc, "Investor Pitch Memo (Draft)")
    add_meta_row(
        doc,
        [
            ("Company", "ReferralMe"),
            ("Date", date.today().isoformat()),
            ("Contact", "amit@referralme.in"),
        ],
    )

    doc.add_heading("One-liner", level=1)
    doc.add_paragraph(
        "ReferralMe is a career-outcomes platform for India that combines structured job search, ethical referral support, and mentorship (AI + human) to increase interview rates and reduce time-to-first-interview.",
    )

    doc.add_heading("Why Now", level=1)
    add_bullets(
        doc,
        [
            "Job search is increasingly noisy: cold DMs and generic applications don’t convert.",
            "Candidates want measurable progress, not just content.",
            "AI lowers the cost of personalized guidance; mentorship converts when users need higher-touch help.",
        ],
    )

    doc.add_heading("Product", level=1)
    add_bullets(
        doc,
        [
            "AI Mentor intake → generates a 7-day plan, outreach templates, and checkpoints.",
            "ATS + application quality tools to improve shortlisting probability.",
            "Referral workflow: seekers request; referrers review and update status.",
            "Human mentorship: paid sessions, verified booking, and post-session feedback loop.",
            "Campus ambassador program to distribute in colleges with tasks, proof review, and rewards.",
        ],
    )

    doc.add_heading("Business Model", level=1)
    add_bullets(
        doc,
        [
            "Mentorship sessions: platform-collected payments; platform fee + payouts to mentors after completion (Phase 2).",
            "Subscription: AI guidance + ATS + weekly execution system (launch after retention proof).",
            "Long-term: employer partnerships and campus program sponsorships.",
        ],
    )

    doc.add_heading("Go-to-Market", level=1)
    add_bullets(
        doc,
        [
            "Campus ambassadors: weekly missions + rewards to drive awareness and signups.",
            "SEO pages for company/role intent (e.g., ‘how to get referral at X’).",
            "Creator partnerships for career guidance distribution.",
            "Referral loops inside the product (invite rewards) to reduce CAC.",
        ],
    )

    doc.add_heading("Moat / Defensibility", level=1)
    add_bullets(
        doc,
        [
            "Workflow + data: repeated user progress data (plans → execution → outcomes) compounds into better guidance.",
            "Supply: curated referrers/mentors + reputation system + program incentives.",
            "Distribution: campus network density + ambassador operations engine.",
        ],
    )

    doc.add_heading("Milestones (Next 2 Quarters)", level=1)
    add_bullets(
        doc,
        [
            "Ship payout ledger + mentor payouts (Phase 2).",
            "Improve AI mentor reliability via paid quota + caching; keep offline mode for resilience.",
            "Launch subscription packaging (₹999/mo) with clear value: plan + ATS + unlimited guidance (within policy).",
            "Add mentor discovery filters and ratings; standardize session outcomes (resume, mock interview, negotiation).",
        ],
    )

    doc.add_heading("What We’re Raising (Placeholder)", level=1)
    doc.add_paragraph(
        "Use this section to add the specific raise amount, runway, and allocation (product, growth, operations).",
    )

    doc.add_heading("Appendix: Visuals", level=1)
    doc.add_paragraph("Sample visuals from the product’s reward and brand assets:")
    if ASSETS["reward_certificate"].exists():
        p = doc.add_paragraph()
        p.add_run().add_picture(str(ASSETS["reward_certificate"]), width=Inches(6.0))

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def main():
    out_dir = ROOT / "docs" / "deliverables"
    product_overview_doc(out_dir / "ReferralMe_Product_Overview.docx")
    investor_pitch_doc(out_dir / "ReferralMe_Investor_Pitch_Memo.docx")
    print("Wrote:", out_dir)


if __name__ == "__main__":
    main()

